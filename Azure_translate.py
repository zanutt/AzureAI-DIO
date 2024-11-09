import requests
from docx import Document
import os

subscription_key = "SUBSCRIPTION_KEY"
endpoint = 'https://api.cognitive.microsofttranslator.com'
location = 'eastus'
target_language = 'pt-br'


#Translate a phrase
def translator_text(text, target_language):
  path ='/translate'
  constructed_url = endpoint + path
  headers={
     'Ocp-Apim-Subscription-Key': subscription_key,
     'Ocp-Apim-Subscription-Region':  location,
     'Content-type': 'application/json',
     'X-ClientTraceId' : str(os.urandom(16))
  }
  body = [{
      'text':text
  }]
  params = {
      'api-version' : '3.0',
      'from' : 'en',
      'to': target_language
  }
  request = requests.post(constructed_url, params=params, headers=headers, json=body)
  response = request.json()
  return response[0]["translations"][0]["text"]

#Translate the intire document
def translate_document(path):
    doc = Document(path)
    full_text =[]
    for paragraph in doc.paragraphs:
        translated_text = translator_text(paragraph.text, target_language)
        full_text.append(translated_text)

    translated_doc = Document()   
    for line in full_text:
      translated_doc.add_paragraph(line)
    path_translated = path.replace('.docx', f'_{target_language}.docx')
    translated_doc.save(path_translated)
    return path_translated